import streamlit as st
import google.generativeai as genai
import os
import time
import docx  # Requer: pip install python-docx
from io import BytesIO
import re
import json
from datetime import datetime
from pymongo import MongoClient
import pandas as pd
import urllib.parse

# Importações para PDF (Novo Módulo)
from reportlab.lib.pagesizes import landscape, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit

# --- 1. CONFIGURAÇÃO DE SEGURANÇA E PROTOCOLO ---
st.set_page_config(
    page_title="TechnoBolt IA - Elite Hub de Governança",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. CONEXÃO MONGODB (RENDER CONFIG) ---
@st.cache_resource
def iniciar_conexao():
    try:
        user = os.environ.get("MONGO_USER", "technobolt")
        password_raw = os.environ.get("MONGO_PASS", "tech@132")
        host = os.environ.get("MONGO_HOST", "cluster0.zbjsvk6.mongodb.net")
        
        password = urllib.parse.quote_plus(password_raw)
        uri = f"mongodb+srv://{user}:{password}@{host}/?appName=Cluster0"
        
        client = MongoClient(uri, serverSelectionTimeoutMS=5000, tlsAllowInvalidCertificates=True)
        client.admin.command('ping') 
        return client['technobolthub']
    except Exception as e:
        st.error(f"⚠️ Erro de conexão com o MongoDB Atlas: {e}")
        return None

db = iniciar_conexao()

# --- 3. GESTÃO DE ESTADO (DNA CORPORATIVO) ---
chaves_sessao = {
    'logged_in': False,
    'user_atual': None,
    'user_plan': 'Standard',
    'is_admin': False,
    'perfil_cliente': {
        "nome_empresa": "TechnoBolt Solutions",
        "setor": "Tecnologia e Consultoria",
        "missao": "Prover governança cognitiva de elite através de IA.",
        "valores": "Inovação, Ética, Precisão, Resiliência.",
        "tom_voz": "Executivo, Autoritário e Analítico"
    },
    'mostrar_resultado': False,
    'resultado_ia': "",
    'titulo_resultado': ""
}

for chave, valor in chaves_sessao.items():
    if chave not in st.session_state:
        st.session_state[chave] = valor

# --- 4. DESIGN SYSTEM (ELITE CORPORATE UI - AJUSTADO) ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    
    /* Global Dark Background */
    html, body, [data-testid="stAppViewContainer"], .stApp {
        background-color: #000000 !important;
        color: #ffffff !important;
        font-family: 'Inter', sans-serif !important;
    }

    /* Sidebar Fix */
    [data-testid="stSidebar"] { background-color: #0a0a0a !important; border-right: 1px solid #222; }
    [data-testid="stHeader"] { background-color: rgba(0,0,0,0) !important; color: white !important; }

    /* Forçar todos os textos para branco */
    p, h1, h2, h3, h4, span, label, div, [data-testid="stMarkdownContainer"] p { color: #ffffff !important; }

    /* FORMS: Corrigir fundo branco */
    [data-testid="stForm"] {
        background-color: #111111 !important; 
        border: 1px solid #333 !important;
        border-radius: 12px !important; 
        padding: 25px !important;
    }

    /* INPUTS & SELECTS: Corrigir fundo branco em todos os estados */
    .stTextInput input, .stTextArea textarea, [data-baseweb="select"] > div, .stSelectbox div {
        background-color: #1a1a1a !important; 
        border: 1px solid #444 !important;
        border-radius: 8px !important; 
        color: #ffffff !important;
    }
    
    /* Foco nos inputs (evitar branco ao clicar) */
    .stTextInput input:focus, .stTextArea textarea:focus {
        background-color: #1a1a1a !important;
        border-color: #666 !important;
        color: #ffffff !important;
    }

    /* BUTTONS: Corrigir botão ficando branco fora do hover e no active */
    .stButton > button {
        width: 100% !important; 
        border-radius: 8px !important; 
        height: 3.2em !important;
        font-weight: 700 !important; 
        background-color: #222222 !important; /* Cor sólida quando parado */
        color: #ffffff !important; 
        border: 1px solid #444 !important; 
        transition: 0.3s ease-in-out !important;
    }

    .stButton > button:hover { 
        background-color: #333333 !important; 
        border-color: #666 !important;
        color: #ffffff !important;
    }
    
    .stButton > button:active, .stButton > button:focus {
        background-color: #444444 !important;
        color: #ffffff !important;
        border-color: #888 !important;
    }

    /* Cards de Resultado */
    .main-card {
        background-color: #111111 !important; 
        border: 1px solid #333 !important;
        border-radius: 12px !important; 
        padding: 25px !important;
    }

    .hero-title { 
        font-size: 32px; font-weight: 800; text-align: center;
        background: linear-gradient(135deg, #ffffff 0%, #444444 100%); 
        -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 25px;
    }
    
    footer { visibility: hidden !important; }
</style>
""", unsafe_allow_html=True)

# --- 5. UTILITÁRIOS E PERSISTÊNCIA ---
def limpar_formatacao(texto):
    texto = texto.replace('**', '').replace('###', '').replace('##', '').replace('#', '')
    texto = re.sub(r'```json.*?```', '', texto, flags=re.DOTALL)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto.strip()

def persistir_interacao(modulo, input_data, output_text, kpis_json):
    if db is not None:
        try:
            log = {
                "usuario": st.session_state.user_atual,
                "timestamp": datetime.now(),
                "modulo": modulo,
                "input": str(input_data)[:500],
                "output": output_text[:1000],
                "kpis": kpis_json
            }
            db["governanca_logs"].insert_one(log)
        except Exception as e:
            print(f"Erro ao salvar log: {e}")

def validar_usuario(username):
    if not username: return False
    return not bool(re.search(r'[\s@]', username))

def gerar_pdf_apresentacao(dados_slides, tema_visual):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    
    # Definição de Cores baseada no Tema
    if "claro" in tema_visual.lower() or "light" in tema_visual.lower() or "branco" in tema_visual.lower():
        bg_color = colors.white
        text_color = colors.black
        accent_color = colors.darkblue
    else: # Padrão Dark (TechnoBolt)
        bg_color = colors.black
        text_color = colors.white
        accent_color = colors.grey
        
    for slide in dados_slides:
        # Fundo
        c.setFillColor(bg_color)
        c.rect(0, 0, width, height, fill=1)
        
        # Header / Título
        c.setFillColor(accent_color)
        c.rect(0, height - 100, width, 100, fill=1)
        
        c.setFillColor(colors.white if bg_color == colors.black else colors.white)
        c.setFont("Helvetica-Bold", 28)
        c.drawCentredString(width / 2, height - 65, slide.get('titulo', 'Sem Título'))
        
        # Conteúdo (Bullets)
        c.setFillColor(text_color)
        c.setFont("Helvetica", 18)
        y_position = height - 150
        
        pontos = slide.get('pontos', [])
        if isinstance(pontos, str): pontos = [pontos]
        
        for ponto in pontos:
            # Quebra de linha simples
            wrapped_text = simpleSplit(f"• {ponto}", "Helvetica", 18, width - 100)
            for line in wrapped_text:
                c.drawString(50, y_position, line)
                y_position -= 25
                if y_position < 50: break # Evita sair da página
            y_position -= 10 # Espaço extra entre itens
            
        # Rodapé
        c.setFillColor(colors.grey)
        c.setFont("Helvetica", 10)
        c.drawString(width - 150, 20, "TechnoBolt Elite Slides")
        
        c.showPage()
        
    c.save()
    buffer.seek(0)
    return buffer

# --- 6. MOTOR DE IA ---
MODEL_FAILOVER_LIST = ["models/gemini-3-flash-preview", "models/gemini-2.5-flash", "models/gemini-2.0-flash", "models/gemini-flash-latest"]

def call_technobolt_ai(prompt, attachments=None, system_context="default"):
    chaves = [os.environ.get(f"GEMINI_CHAVE_{i}") for i in range(1, 8)]
    chaves = [k for k in chaves if k] or [os.environ.get("GEMINI_API_KEY")]

    p = st.session_state.perfil_cliente
    dna_context = f"DNA: {p['nome_empresa']}. Tom: {p['tom_voz']}.\n"
    
    # Instrução padrão para módulos analíticos (mantém JSON de KPIs)
    kpi_instruction = (
        "\nOBRIGATÓRIO: No final da resposta, adicione EXATAMENTE um bloco JSON estruturado entre ```json e ``` "
        "com as chaves: 'faturamento', 'margem', 'riscos_count', 'prazos_alerta'."
    )
    
    # Instrução especial para Slides
    if system_context == "slides":
        sys_instr = (
            "Você é um Designer de Apresentações Executivas. "
            "Gere o conteúdo em formato JSON ESTRUTURADO. "
            "O formato deve ser uma lista de objetos: [{'titulo': '...', 'pontos': ['...', '...']}]. "
            "Crie entre 4 a 6 slides. Seja direto e use tópicos curtos."
        )
    else:
        contexts = {
            "mckinsey": "Persona: Sócio McKinsey. Framework: 7S e MECE.",
            "email_intel": "Persona: CCO. Triagem diplomática.",
            "briefing": "Persona: Diretor de Inteligência.",
            "ata": "Persona: Secretário de Governança B3.",
            "churn": "Persona: Especialista em Retenção.",
            "master": "Persona: COO. Consolidação semanal.",
            "default": "Consultoria Sênior TechnoBolt."
        }
        sys_instr = dna_context + contexts.get(system_context, contexts["default"]) + kpi_instruction
    
    for key in chaves:
        try:
            genai.configure(api_key=key)
            for model_name in MODEL_FAILOVER_LIST:
                try:
                    model = genai.GenerativeModel(model_name, system_instruction=sys_instr)
                    payload = [prompt] + attachments if attachments else prompt
                    response = model.generate_content(payload)
                    full_text = response.text
                    
                    if system_context == "slides":
                        return full_text, f"{model_name.split('/')[-1]}"
                    
                    # Lógica padrão para outros módulos
                    kpis = {"faturamento": 0, "margem": 0, "riscos_count": 0, "prazos_alerta": 0}
                    json_match = re.search(r'```json\n(.*?)\n```', full_text, re.DOTALL)
                    if json_match:
                        try: kpis = json.loads(json_match.group(1))
                        except: pass
                    clean_res = limpar_formatacao(full_text)
                    persistir_interacao(system_context, prompt, clean_res, kpis)
                    return clean_res, f"{model_name.split('/')[-1]}"
                except: continue
        except: continue
    return "⚠️ Motores em manutenção.", "OFFLINE"

# --- 7. AUTENTICAÇÃO ---
if not st.session_state.logged_in:
    tab1, tab2 = st.tabs(["Hub Acesso", "Solicitar Acesso"])
    with tab1:
        with st.form("auth_hub"):
            st.markdown("<h1 class='hero-title'>TECHNOBOLT HUB</h1>", unsafe_allow_html=True)
            u = st.text_input("Operador ID", placeholder="ID")
            k = st.text_input("Chave PIN", type="password")
            
            if st.form_submit_button("CONECTAR"):
                if db is not None:
                    try:
                        user_db = db["usuarios"].find_one({"usuario": u, "senha": k})
                        if user_db:
                            if user_db.get("status") == "ativo":
                                st.session_state.logged_in = True
                                st.session_state.user_atual = u
                                st.session_state.user_plan = user_db.get("plano", "Standard")
                                st.session_state.is_admin = user_db.get("is_admin", False)
                                st.rerun()
                            else:
                                st.warning("Conta aguardando ativação administrativa.")
                        else:
                            st.error("ID ou PIN incorretos.")
                    except Exception as e:
                        st.error(f"Erro ao consultar banco de dados: {e}")
                else:
                    st.error("SISTEMA OFFLINE: Conexão com MongoDB não estabelecida.")
                    
    with tab2:
        with st.form("request_access"):
            new_u = st.text_input("ID Desejado (Sem espaços ou @)")
            new_k = st.text_input("PIN de Segurança", type="password")
            plan_req = st.selectbox("Plano Desejado", ["Standard", "Advanced", "Executive"])
            
            if st.form_submit_button("SOLICITAR"):
                if db is not None:
                    if validar_usuario(new_u):
                        if not db["usuarios"].find_one({"usuario": new_u}):
                            db["usuarios"].insert_one({
                                "usuario": new_u, "senha": new_k, 
                                "plano": plan_req, "status": "inativo",
                                "is_admin": False,
                                "criado_em": datetime.now()
                            })
                            st.success("Solicitação em processamento.")
                        else: st.error("Este ID já está registrado.")
                    else: st.error("O ID não pode conter espaços ou o caractere '@'.")
                else:
                    st.error("Não é possível solicitar acesso com o banco offline.")
    st.stop()

# --- 8. NAVEGAÇÃO E PAYWALL ---
with st.sidebar:
    st.markdown(f"<h2 style='color:#ffffff; text-align:center;'>TECHNOBOLT</h2>", unsafe_allow_html=True)
    st.markdown(f"<p style='text-align:center; font-size:10px;'>MODO: {st.session_state.user_plan.upper()}</p>", unsafe_allow_html=True)
    st.markdown("---")
    
    opcoes = ["Centro de Comando", "Analisador de Documentos", "Analisador de E-mails", "Gerador de Emails", "Briefing Estratégico", "Gerador de Atas", "Mercado & Churn", "Relatório Semanal", "Criar Apresentação"]
    
    # AJUSTE: Permite acesso à gestão para ID "admin" ou qualquer um com flag is_admin
    if st.session_state.user_atual == "admin" or st.session_state.is_admin:
        opcoes.append("Gestão de Acesso")
        
    escolha = st.radio("NAV", opcoes, label_visibility="collapsed")
    st.markdown("<div style='height: 10vh;'></div>", unsafe_allow_html=True)
    if st.button("🚪 Sair"):
        st.session_state.logged_in = False
        st.rerun()

restritos_standard = ["Analisador de Documentos", "Analisador de E-mails", "Mercado & Churn", "Criar Apresentação"]
if st.session_state.user_plan == "Standard" and escolha in restritos_standard:
    st.markdown(f"<div class='main-card'><h2>🚀 Upgrade Necessário</h2><p>O módulo <b>{escolha}</b> está disponível apenas para parceiros Advanced e Executive.</p></div>", unsafe_allow_html=True)
    st.stop()

# --- 9. MÓDULOS ---

if escolha == "Centro de Comando":
    st.markdown("<h1 class='hero-title'>Dashboard Cognitivo</h1>", unsafe_allow_html=True)
    
    if db is not None:
        logs = list(db["governanca_logs"].find({"usuario": st.session_state.user_atual}).sort("timestamp", -1).limit(10))
        if logs:
            c1, c2, c3 = st.columns(3)
            riscos = sum([l.get("kpis", {}).get("riscos_count", 0) for l in logs])
            alertas = sum([l.get("kpis", {}).get("prazos_alerta", 0) for l in logs])
            c1.metric("Riscos Identificados", riscos)
            c2.metric("Prazos Críticos", alertas)
            c3.metric("Ações Recentes", len(logs))
            df = pd.DataFrame([{"Data": l["timestamp"], "Riscos": l.get("kpis", {}).get("riscos_count", 0)} for l in logs])
            st.line_chart(df.set_index("Data"))
        else:
            st.info("Execute análises para popular o dashboard estratégico.")
    else:
        st.warning("Métricas indisponíveis em modo offline.")

elif escolha == "Criar Apresentação":
    st.markdown("<div class='main-card'><h2>Gerador de Slides Executivos</h2></div>", unsafe_allow_html=True)
    
    with st.form("form_slides"):
        tema_slides = st.text_area("Descreva o conteúdo e objetivo da apresentação:")
        estilo_visual = st.selectbox("Estilo Visual / Tema:", 
                                     ["TechnoBolt Dark (Padrão)", "Minimalista Claro", "Corporativo Azul", "High Tech Neon"])
        
        if st.form_submit_button("GERAR APRESENTAÇÃO"):
            if tema_slides.strip():
                with st.spinner("Desenhando slides e estruturando narrativa..."):
                    prompt = f"Crie uma apresentação sobre: {tema_slides}. O tema visual é: {estilo_visual}. Retorne APENAS o JSON."
                    raw_res, mot = call_technobolt_ai(prompt, None, "slides")
                    
                    # Tentativa de parse do JSON
                    json_match = re.search(r'```json\n(.*?)\n```', raw_res, re.DOTALL)
                    dados_slides = []
                    
                    if json_match:
                        try: 
                            dados_slides = json.loads(json_match.group(1))
                        except: pass
                    else:
                        # Fallback simples se o JSON falhar
                        try: dados_slides = json.loads(raw_res)
                        except: pass
                        
                    if dados_slides:
                        # Gerar PDF
                        pdf_buffer = gerar_pdf_apresentacao(dados_slides, estilo_visual)
                        
                        st.session_state.resultado_ia = "Apresentação gerada com sucesso. Clique abaixo para baixar."
                        st.session_state.titulo_resultado = f"Slides: {tema_slides[:30]}..."
                        st.session_state.mostrar_resultado = False # Não mostra texto, mostra botão direto
                        
                        st.success("Slides Renderizados!")
                        st.download_button(
                            label="📥 BAIXAR PDF (PAISAGEM)",
                            data=pdf_buffer,
                            file_name="apresentacao_technobolt.pdf",
                            mime="application/pdf"
                        )
                        
                        # Preview rápido dos tópicos
                        with st.expander("Visualizar Estrutura Gerada"):
                            for s in dados_slides:
                                st.markdown(f"**{s.get('titulo')}**")
                                for p in s.get('pontos', []):
                                    st.markdown(f"- {p}")
                    else:
                        st.error("Erro ao estruturar os dados da apresentação. Tente novamente.")
            else:
                st.warning("Descreva o tema da apresentação.")

elif escolha == "Analisador de Documentos":
    st.markdown("<div class='main-card'><h2>Analisador de Documentos</h2></div>", unsafe_allow_html=True)
    with st.form("form_docs"):
        up = st.file_uploader("Submeter PDF/DOCX", type=['pdf', 'docx'])
        if st.form_submit_button("EXECUTAR PROTOCOLO"):
            if up:
                with st.spinner("Auditando..."):
                    content = up.read()
                    res, mot = call_technobolt_ai("Audite este documento e gere KPIs.", [content], "mckinsey")
                    st.session_state.resultado_ia = res
                    st.session_state.titulo_resultado = f"Auditoria McKinsey ({mot})"
                    st.session_state.mostrar_resultado = True
                    st.rerun()

elif escolha == "Analisador de E-mails":
    st.markdown("<div class='main-card'><h2>Analisador de E-mails</h2></div>", unsafe_allow_html=True)
    with st.form("form_emails"):
        lote = st.text_area("Cole aqui os blocos de e-mail para triagem:", height=250)
        if st.form_submit_button("EXECUTAR TRIAGEM"):
            if lote.strip():
                with st.spinner("CCO analisando comunicações..."):
                    res, mot = call_technobolt_ai(lote, None, "email_intel")
                    st.session_state.resultado_ia = res
                    st.session_state.titulo_resultado = f"Triagem Executiva ({mot})"
                    st.session_state.mostrar_resultado = True
                    st.rerun()
            else:
                st.warning("⚠️ Forneça o conteúdo dos e-mails para processamento.")

    if db is not None:
        st.markdown("### 📥 Histórico de Triagens Recentes")
        try:
            hist = list(db["governanca_logs"].find({
                "modulo": "email_intel", 
                "usuario": st.session_state.user_atual
            }).sort("timestamp", -1).limit(3))
            
            if hist:
                for h in hist:
                    data_formatada = h['timestamp'].strftime('%d/%m/%Y %H:%M')
                    with st.expander(f"Triagem realizada em {data_formatada}"):
                        st.write(h['output'])
            else:
                st.info("Nenhuma triagem encontrada no histórico.")
        except Exception as e:
            st.error(f"Erro ao carregar histórico: {e}")

elif escolha == "Gerador de Emails":
    st.markdown("<div class='main-card'><h2>Gerador de Emails</h2></div>", unsafe_allow_html=True)
    with st.form("form_gen_mail"):
        pauta = st.text_area("Descreva o assunto ou pauta do e-mail:")
        if st.form_submit_button("GERAR EMAIL"):
            with st.spinner("Redigindo e-mail diplomático..."):
                res, mot = call_technobolt_ai(f"Gere um email profissional sobre: {pauta}", None, "default")
                st.session_state.resultado_ia = res
                st.session_state.titulo_resultado = f"Email Redigido ({mot})"
                st.session_state.mostrar_resultado = True
                st.rerun()
    
    if db is not None:
        st.markdown("### 🕒 Últimas comunicações")
        hist = list(db["governanca_logs"].find({"modulo": "default", "usuario": st.session_state.user_atual}).sort("timestamp", -1).limit(3))
        for h in hist:
            with st.expander(f"E-mail de {h['timestamp'].strftime('%d/%m - %H:%M')}"):
                st.write(h['output'])

elif escolha == "Briefing Estratégico":
    st.markdown("<div class='main-card'><h2>Briefing Estratégico</h2></div>", unsafe_allow_html=True)
    with st.form("form_briefing"):
        alvo = st.text_input("Empresa ou Setor para Análise:")
        if st.form_submit_button("GERAR SCAN"):
            with st.spinner("Coletando inteligência..."):
                res, mot = call_technobolt_ai(f"Gere um briefing estratégico sobre: {alvo}", None, "briefing")
                st.session_state.resultado_ia = res
                st.session_state.titulo_resultado = f"Scan de Mercado: {alvo} ({mot})"
                st.session_state.mostrar_resultado = True
                st.rerun()

elif escolha == "Gerador de Atas":
    st.markdown("<div class='main-card'><h2>Gerador de Atas</h2></div>", unsafe_allow_html=True)
    with st.form("form_atas"):
        notas = st.text_area("Notas ou Transcrições da Reunião:", height=250)
        if st.form_submit_button("FORMALIZAR ATA"):
            with st.spinner("Estruturando Governança..."):
                res, mot = call_technobolt_ai(notas, None, "ata")
                st.session_state.resultado_ia = res
                st.session_state.titulo_resultado = f"Ata Formalizada ({mot})"
                st.session_state.mostrar_resultado = True
                st.rerun()

elif escolha == "Mercado & Churn":
    st.markdown("<div class='main-card'><h2>Inteligência de Mercado & Churn</h2></div>", unsafe_allow_html=True)
    with st.form("form_churn"):
        dados_cli = st.text_area("Feedbacks ou métricas de comportamento do cliente:")
        if st.form_submit_button("CALCULAR RISCO"):
            with st.spinner("Avaliando retenção..."):
                res, mot = call_technobolt_ai(dados_cli, None, "churn")
                st.session_state.resultado_ia = res
                st.session_state.titulo_resultado = f"Análise de Retenção ({mot})"
                st.session_state.mostrar_resultado = True
                st.rerun()

elif escolha == "Relatório Semanal":
    st.markdown("<div class='main-card'><h2>Relatório Semanal de KPIs</h2></div>", unsafe_allow_html=True)
    with st.form("form_semanal"):
        kpis = st.text_area("Fatos, métricas e marcos da semana:")
        if st.form_submit_button("CONSOLIDAR RELATÓRIO"):
            with st.spinner("COO consolidando dados..."):
                res, mot = call_technobolt_ai(kpis, None, "master")
                st.session_state.resultado_ia = res
                st.session_state.titulo_resultado = f"Relatório Semanal ({mot})"
                st.session_state.mostrar_resultado = True
                st.rerun()

elif escolha == "Gestão de Acesso" and (st.session_state.user_atual == "admin" or st.session_state.is_admin):
    st.markdown("## 🔐 Governança de Acessos")
    if db is not None:
        usuarios_lista = list(db["usuarios"].find())
        for u_data in usuarios_lista:
            user_nome = u_data.get('usuario', 'Usuário Sem Nome')
            user_plano = u_data.get('plano', 'Standard')
            user_status = u_data.get('status', 'inativo')
            is_user_admin = u_data.get('is_admin', False)
            perfil_md = f"### Operador: `{user_nome}` {' [ADMIN]' if is_user_admin else ''}\n---\n- **Plano Atual:** {user_plano}\n- **Status:** {user_status}"
            st.markdown(perfil_md)
            col_p, col_s, col_adm = st.columns([1, 1, 1])
            with col_p:
                lista_planos = ["Standard", "Advanced", "Executive"]
                try: idx_plano = lista_planos.index(user_plano)
                except ValueError: idx_plano = 0
                new_p = st.selectbox("Mudar Plano", lista_planos, index=idx_plano, key=f"plan_{user_nome}")
            with col_s:
                lista_status = ["ativo", "inativo"]
                idx_status = 0 if user_status == "ativo" else 1
                new_s = st.selectbox("Mudar Status", lista_status, index=idx_status, key=f"stat_{user_nome}")
            with col_adm:
                st.markdown("<br>", unsafe_allow_html=True)
                if not is_user_admin:
                    if st.button("Promover a Administrador", key=f"make_adm_{user_nome}"):
                        db["usuarios"].update_one({"usuario": user_nome}, {"$set": {"is_admin": True}})
                        st.success(f"{user_nome} agora é Admin.")
                        st.rerun()
                else: st.info("Acesso Admin")
            if st.button("Aplicar Alterações", key=f"save_{user_nome}"):
                db["usuarios"].update_one({"usuario": user_nome}, {"$set": {"plano": new_p, "status": new_s}})
                st.success(f"Dossiê de {user_nome} atualizado.")
                st.rerun()
            st.markdown("---")
    else: st.error("Banco de dados inacessível.")

# --- 10. EXIBIÇÃO DE RESULTADOS ---
if st.session_state.mostrar_resultado:
    st.markdown("---")
    _, mid, _ = st.columns([1, 8, 1])
    with mid:
        texto_limpo = st.session_state.resultado_ia
        st.markdown(f"<div class='result-card-elite'><h2 style='color:#ffffff !important;'>{st.session_state.titulo_resultado}</h2><div style='color:#eee !important; white-space: pre-wrap; font-size: 15px;'>{texto_limpo}</div></div>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        if c1.button("📥 BAIXAR DOCX"):
            doc = docx.Document()
            doc.add_heading(st.session_state.titulo_resultado, 0)
            doc.add_paragraph(texto_limpo)
            buf = BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button("Confirmação de Download", buf, "relatorio_technobolt.docx")
        if c2.button("✖️ FECHAR"):
            st.session_state.mostrar_resultado = False
            st.rerun()

st.markdown("<br><br>", unsafe_allow_html=True)
st.caption(f"TechnoBolt Solutions © 2026 | Hub Elite v2.0")
